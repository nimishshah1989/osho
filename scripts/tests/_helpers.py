"""Test helpers shared across the ingester and the operational-pipeline suites.

Kept out of `conftest.py` so it's importable as a normal module: pytest
treats `conftest.py` as a fixture-and-hook plugin and doesn't make it
available via `from conftest import …`.
"""
from __future__ import annotations


def make_docx(
    path,
    *,
    title: str = "Sample Discourse ~ 01",
    language: str = "EN",
    translated_from: str = "none",
    time: str = "1987-03-08-xm",
    body=None,
):
    """Build a small `.docx` with @-field headers and styled body paragraphs.

    Both the ingester tests and the operational-pipeline tests exercise the
    exact same Word format via this helper.

    Args:
        path: where to write the file (any os.PathLike).
        title / language / translated_from / time: values for the matching
            `@field=` headers.
        body: list of body entries. Each entry can be:
            - a plain `str` (rendered with style "ctp - Osho Talking"), or
            - a `(text, style)` tuple to override the style.
            Pass `[]` for an empty body (e.g. Delete-flow files). Pass
            `None` for a default one-paragraph body.
    """
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    for name in (
        "ctp - Event Info",
        "ctp - Osho Talking",
        "ctp - Other Talking 1",
        "ctp - Sutra/Question",
        "ctp - Poem",
    ):
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    for line in (
        f"@title={title}",
        f"@language={language}",
        f"@translatedFrom={translated_from}",
        f"@time={time}",
        "@eventText=",
    ):
        doc.add_paragraph(line, style="ctp - Event Info")

    if body is None:
        body = [("A single paragraph of body text.", "ctp - Osho Talking")]
    for entry in body:
        if isinstance(entry, str):
            doc.add_paragraph(entry, style="ctp - Osho Talking")
        else:
            text, style = entry
            if style is None:
                # Pass through to python-docx's Normal style → ingester
                # produces role=None, which is the "plain body text" path.
                doc.add_paragraph(text)
            else:
                doc.add_paragraph(text, style=style)
    doc.save(str(path))
