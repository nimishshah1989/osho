"""Ingester round-trips Word `ctp - …` paragraph styles into the
`paragraphs.role` column, and the API surfaces them on /discourse and /search.

These tests build a real .docx in-memory using python-docx so we cover the
full path from style-name → role → DB → JSON, not just the regex normaliser.
"""

import os
import sqlite3
import sys
import tempfile

import pytest

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

from scripts.ingest_docx import (  # noqa: E402
    _normalise_role,
    parse_docx,
    upsert,
    _ensure_role_column,
    _ensure_translated_from_column,
)


# ── Style → role normalisation ────────────────────────────────────────────

@pytest.mark.parametrize(
    "style_name, expected",
    [
        ("ctp - Osho Talking", "osho_talking"),
        ("ctp - Other Talking 1", "other_talking_1"),
        ("ctp - Other Talking 2", "other_talking_2"),
        ("ctp - Sutra/Question", "sutra_question"),
        ("ctp - Poem", "poem"),
        ("ctp - Comments", "comments"),
        ("ctp - Short comments", "short_comments"),
        ("ctp - Our translation", "our_translation"),
        ("ctp - Notes", "notes"),
        ("ctp - Title", "title"),
        ("ctp - Event Info", "event_info"),
        # case-insensitive on the prefix
        ("CTP - Osho Talking", "osho_talking"),
        # whitespace tolerant
        ("ctp -   Osho   Talking", "osho_talking"),
        # Punctuation in the label is collapsed into a single underscore so the
        # slug is always a safe map key / CSS class fragment. Guards against
        # future style names like "ctp - Q & A" leaking ampersands into the role.
        ("ctp - Q & A", "q_a"),
        ("ctp - Footnote (1)", "footnote_1"),
        ("ctp - !!!", None),  # nothing identifier-like left → None
        # Non-ctp styles → None (treated as plain body text)
        ("Normal", None),
        ("Heading 1", None),
        ("", None),
        (None, None),
    ],
)
def test_normalise_role(style_name, expected):
    assert _normalise_role(style_name) == expected


# ── End-to-end: docx → DB → role round-trips ──────────────────────────────

def _make_docx(path: str) -> None:
    """Build a small .docx with @-headers and a mix of ctp- styles."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # Define the styles we'll use so python-docx doesn't fall back to Normal.
    for name in (
        "ctp - Event Info",
        "ctp - Osho Talking",
        "ctp - Other Talking 1",
        "ctp - Sutra/Question",
        "ctp - Poem",
    ):
        if name not in [s.name for s in doc.styles]:
            from docx.enum.style import WD_STYLE_TYPE
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.font.size = Pt(12)

    # Header lines — Event Info style, exactly the convention Antar uses
    for line in (
        "@title=Sample Discourse ~ 01",
        "@language=EN",
        "@translatedFrom=none",
        "@time=1987-03-08-xm",
        "@eventText=",
    ):
        p = doc.add_paragraph(line, style="ctp - Event Info")

    # Body — mixed roles
    doc.add_paragraph(
        "INTERVIEWER: Why are you saying this?", style="ctp - Other Talking 1"
    )
    doc.add_paragraph(
        "OSHO: Because the question itself reveals the answer.",
        style="ctp - Osho Talking",
    )
    doc.add_paragraph(
        "Sutra to be commented upon today.", style="ctp - Sutra/Question"
    )
    doc.add_paragraph("A short verse for the moment.", style="ctp - Poem")
    # An untyped paragraph — should become role=None
    doc.add_paragraph("A bare paragraph with no ctp style.")

    doc.save(path)


def _seed_minimal_db(path: str) -> None:
    """Empty DB matching production schema, including the FTS table the
    upsert path expects to populate."""
    conn = sqlite3.connect(path)
    conn.executescript(
        """
        CREATE TABLE events (
            id TEXT PRIMARY KEY, title TEXT NOT NULL,
            date TEXT, location TEXT, language TEXT
        );
        CREATE TABLE paragraphs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_id TEXT NOT NULL,
            sequence_number INTEGER NOT NULL,
            content TEXT NOT NULL,
            is_embedded BOOLEAN DEFAULT 0
        );
        CREATE VIRTUAL TABLE paragraphs_fts USING fts5(
            content, title UNINDEXED, event_id UNINDEXED,
            paragraph_id UNINDEXED, sequence_number UNINDEXED,
            title_search,
            tokenize = 'porter unicode61 remove_diacritics 2'
        );
        """
    )
    conn.commit()
    conn.close()


def test_ingest_round_trips_role_to_db(tmp_path):
    docx_path = tmp_path / "sample.docx"
    db_path = tmp_path / "test.db"
    _make_docx(str(docx_path))
    _seed_minimal_db(str(db_path))

    talk = parse_docx(docx_path)
    assert talk.title == "Sample Discourse ~ 01"
    assert talk.language == "English"
    # Roles parsed off the in-memory paragraphs
    roles = [p.role for p in talk.paragraphs]
    assert roles == [
        "other_talking_1",
        "osho_talking",
        "sutra_question",
        "poem",
        None,
    ]

    with sqlite3.connect(db_path) as conn:
        _ensure_translated_from_column(conn)
        _ensure_role_column(conn)
        upsert(conn, talk)
        rows = conn.execute(
            "SELECT sequence_number, content, role FROM paragraphs"
            " ORDER BY sequence_number"
        ).fetchall()

    assert [r[2] for r in rows] == [
        "other_talking_1",
        "osho_talking",
        "sutra_question",
        "poem",
        None,
    ]
    # Content survived intact too
    assert "Why are you saying this" in rows[0][1]
    assert "the question itself reveals" in rows[1][1]
