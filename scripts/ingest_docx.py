"""Bulk-ingest Osho discourses from `.docx` files into the search database.

Each `.docx` follows the format Antar uses (see the example files in the
project root):

    @title=A New Vision of Women's Liberation ~ 01
    @language=EN
    @translatedFrom=none
    @sourceShort=             (only when @translatedFrom != "none")
    @time=1987-03-08-xm
    @place=                   (optional)
    @eventText=
    <full paragraph 1>

(The legacy `@theme=` header is still accepted for backward compatibility
with Antar's pre-2026-05 documents; new Word files Sugit ships should
omit it — themes are auto-classified from content.)

    <full paragraph 2>

    ...

Filename convention (recommended but not required):
    "Talk Title ~ 01_LHI.docx"   — Hindi
    "Talk Title ~ 01_LEN.docx"   — English

Idempotency
-----------
Records are upserted on the composite key (title, language). Re-running this
script on the same `.docx` file updates the existing event in place instead
of creating a duplicate. The FTS rows for that event are regenerated.

Usage
-----
    python3 scripts/ingest_docx.py /path/to/docs/
    python3 scripts/ingest_docx.py one_file.docx

Run on the server. The backend reads from the same `data/osho.db` that this
script writes to, so changes are visible immediately — no restart required
(the FTS5 virtual table is updated transactionally).
"""
from __future__ import annotations

import argparse
import os
import re
import sqlite3
import sys
import unicodedata
import uuid
from dataclasses import dataclass, field
from pathlib import Path

# Re-use the same Devanagari normalisation as build_fts / cloud_api so the
# anusvara-equivalence promise holds for every ingested talk.
from build_fts import normalize_devanagari


BASE_DIR = Path(__file__).resolve().parent.parent
DB_PATH = BASE_DIR / "data" / "osho.db"

# Whitelisted @-fields we read from the doc header.
# `theme` is legacy (pre-2026-05) — still accepted so old docs ingest cleanly,
# but new files from Sugit don't carry it.
_HEADER_FIELDS = {
    "title", "language", "translatedfrom", "sourceshort",
    "time", "theme", "place", "eventtext",
}

# Match a header line like "@field=value" — the `=` is required; trailing
# whitespace is preserved on the value (we strip it ourselves).
_HEADER_RE = re.compile(r"^@(\w+)\s*=\s*(.*)$")

# Language code from filename suffix _LHI / _LEN / _LZH …
_FILENAME_LANG_RE = re.compile(r"_L([A-Z]{2,3})\.docx$", re.IGNORECASE)

# Two-letter language codes → canonical full names matching the existing DB.
_LANG_MAP = {
    "EN": "English",
    "HI": "Hindi",
}


@dataclass
class Paragraph:
    text: str
    role: str | None = None  # see _normalise_role


@dataclass
class TalkRecord:
    title: str
    language: str
    translated_from: str | None = None
    # Short book title for translations — surfaced in search results so a
    # reader who finds a translated talk can see which published volume
    # it came from. Only meaningful when `translated_from` is not "none".
    source_short: str | None = None
    time: str | None = None
    theme: str | None = None
    place: str | None = None
    paragraphs: list[Paragraph] = field(default_factory=list)
    source_path: str = ""

    @property
    def event_key(self) -> tuple[str, str]:
        """Composite upsert key — same talk in the same language is the same row."""
        return (self.title.strip(), self.language.strip())


# Antar's Word documents tag every paragraph with a `ctp - ...` style that
# carries the semantic role (Osho speaking vs. interviewer vs. sutra vs. poem
# vs. footnote, etc). We strip the prefix and slugify the rest so the role is
# a stable machine identifier the frontend can map to a CSS class. Paragraphs
# without a `ctp -` style get role=None and render as plain body text.
_CTP_STYLE_RE = re.compile(r"^ctp\s*-\s*(.+?)\s*$", re.IGNORECASE)


def _normalise_role(style_name: str | None) -> str | None:
    if not style_name:
        return None
    m = _CTP_STYLE_RE.match(style_name)
    if not m:
        return None
    label = m.group(1).strip().lower()
    # Sugit's Hindi files use the same semantic styles as English but with a
    # trailing language qualifier — "ctp - Osho Talking (Hindi)" carries the
    # same role as "ctp - Osho Talking", differing only in the font the Word
    # template applies. Strip the qualifier so both map to the same slug;
    # otherwise the frontend would render them with two different CSS
    # classes for what is the same semantic paragraph type.
    label = re.sub(r"\s*\((hindi|devanagari|english|en|hi)\)\s*$", "", label)
    # "Sutra/Question" → "sutra_question", "Other Talking 1" → "other_talking_1",
    # "Q & A" → "q_a". `[^\w]+` collapses any run of non-identifier characters
    # so the slug is always a safe map-key / CSS-class fragment.
    slug = re.sub(r"[^\w]+", "_", label).strip("_")
    return slug or None


def _import_python_docx():
    try:
        import docx  # type: ignore
    except ImportError:
        sys.exit(
            "python-docx is not installed. Run: pip install python-docx\n"
            "(or `pip install -r requirements.txt`)"
        )
    return docx


def _read_docx_paragraphs(path: Path) -> list[Paragraph]:
    """Return non-empty paragraphs (text + semantic role) in document order."""
    docx = _import_python_docx()
    doc = docx.Document(str(path))
    out: list[Paragraph] = []
    for p in doc.paragraphs:
        text = p.text
        if not text or not text.strip():
            continue
        style_name = p.style.name if p.style else None
        out.append(Paragraph(text=text, role=_normalise_role(style_name)))
    return out


def _parse_header_and_body(
    paragraphs: list[Paragraph],
) -> tuple[dict[str, str], list[Paragraph]]:
    """Split a list of paragraphs into (header_dict, body_paragraphs).

    The header continues until we hit `@eventText=` (anything after the `=`
    on that line is treated as the first body paragraph, in case the author
    pasted the first paragraph inline). All subsequent paragraphs are body.
    """
    header: dict[str, str] = {}
    body: list[Paragraph] = []
    in_body = False

    for para in paragraphs:
        if in_body:
            body.append(para)
            continue

        m = _HEADER_RE.match(para.text.strip())
        if not m:
            # Header is over the moment we see a non-@ line.
            in_body = True
            body.append(para)
            continue

        key = m.group(1).lower()
        value = m.group(2).rstrip()

        if key == "eventtext":
            in_body = True
            if value:  # rare: content inline on same line as @eventText=
                body.append(Paragraph(text=value, role=para.role))
        elif key in _HEADER_FIELDS:
            header[key] = value
        # silently ignore unknown @-fields so old docs don't error

    return header, body


def _language_from(header: dict[str, str], filename: Path) -> str:
    """Resolve a canonical language name, preferring the @language header.

    Falls back to the `_LXX` suffix on the filename. Codes are mapped to the
    full names already used in the events table (English, Hindi, …).
    """
    raw = header.get("language", "").strip().upper()
    if not raw:
        m = _FILENAME_LANG_RE.search(filename.name)
        raw = m.group(1).upper() if m else ""
    if not raw:
        return "Unknown"
    return _LANG_MAP.get(raw, raw.title())


def parse_docx_header(path: Path) -> tuple[dict[str, str], str]:
    """Parse just the @-field headers from a .docx and return (header, language).

    Used by the Delete flow, where a "delete this record" Word file may be
    empty below the headers (Sugit's convention). `parse_docx` itself
    insists on body paragraphs because it's about ingesting content; for
    deletion we only need the (title, language) identifier."""
    paras = _read_docx_paragraphs(path)
    if not paras:
        raise ValueError(f"{path.name}: empty document (no @title= header found)")
    header, _ = _parse_header_and_body(paras)
    if not header.get("title"):
        raise ValueError(
            f"{path.name}: missing @title= header line "
            f"(found headers: {sorted(header.keys()) or 'none'})"
        )
    return header, _language_from(header, path)


def parse_docx(path: Path) -> TalkRecord:
    """Read a single `.docx` and return a TalkRecord."""
    paras = _read_docx_paragraphs(path)
    if not paras:
        raise ValueError(f"{path.name}: empty document")

    header, body = _parse_header_and_body(paras)
    if not header.get("title"):
        raise ValueError(
            f"{path.name}: missing @title= header line "
            f"(found headers: {sorted(header.keys()) or 'none'})"
        )

    body = [Paragraph(text=p.text.strip(), role=p.role) for p in body if p.text.strip()]
    if not body:
        raise ValueError(f"{path.name}: no body paragraphs after @eventText=")

    translated_from = (header.get("translatedfrom") or "").strip() or None
    # Sugit writes `@translatedFrom=` with the same language codes he uses
    # on `@language=` (EN, HI, …), so map them through the same table.
    # The existing corpus stores the full name ("Hindi" / "English") on the
    # translated_from column; keeping a single canonical form is what
    # makes filters like "originals only" work across old + new records.
    # The literal "none" sentinel and any unrecognised value pass through
    # unchanged so we don't silently rewrite something we didn't expect.
    if translated_from:
        code = translated_from.upper()
        if code in _LANG_MAP:
            translated_from = _LANG_MAP[code]
    source_short = (header.get("sourceshort") or "").strip() or None
    # Sugit's convention: `@sourceShort` only carries meaning for translations.
    # If it appears on an original-language record we drop it rather than
    # storing a value that the UI would mis-display as a book-of-origin.
    if source_short and (not translated_from or translated_from.lower() == "none"):
        print(
            f"  WARN  {path.name}: @sourceShort ignored — record is not a translation",
            file=sys.stderr,
        )
        source_short = None

    return TalkRecord(
        title=header["title"].strip(),
        language=_language_from(header, path),
        translated_from=translated_from,
        source_short=source_short,
        time=(header.get("time") or "").strip() or None,
        theme=(header.get("theme") or "").strip() or None,
        place=(header.get("place") or "").strip() or None,
        paragraphs=body,
        source_path=str(path),
    )


# ─── Database upsert ────────────────────────────────────────────────────────

def _ensure_translated_from_column(conn: sqlite3.Connection) -> None:
    cols = {r[1] for r in conn.execute("PRAGMA table_info(events)").fetchall()}
    if "translated_from" not in cols:
        conn.execute("ALTER TABLE events ADD COLUMN translated_from TEXT")


def _ensure_source_short_column(conn: sqlite3.Connection) -> None:
    cols = {r[1] for r in conn.execute("PRAGMA table_info(events)").fetchall()}
    if "source_short" not in cols:
        conn.execute("ALTER TABLE events ADD COLUMN source_short TEXT")


def _ensure_role_column(conn: sqlite3.Connection) -> None:
    cols = {r[1] for r in conn.execute("PRAGMA table_info(paragraphs)").fetchall()}
    if "role" not in cols:
        conn.execute("ALTER TABLE paragraphs ADD COLUMN role TEXT")


# Characters that all stand in for the " ~ NN " series-part separator in
# titles. Antar's Word files, Sugit's batches and the legacy corpus type this
# inconsistently — plain ASCII tilde, the swung dash, the fullwidth tilde, the
# whole hyphen/dash family and the minus sign all show up. They are collapsed
# to a single character for *matching only* so two otherwise-identical titles
# that differ solely in which glyph was typed resolve to the same record.
_TITLE_SEPARATORS = "~∼～⁓‐‑‒–—―−-"
_TITLE_SEP_RE = re.compile(f"[{re.escape(_TITLE_SEPARATORS)}]")
_TITLE_WS_RE = re.compile(r"\s+")


def _canonical_title(title: str) -> str:
    """Normalisation-tolerant key for matching a title to an existing record.

    Two titles that look identical but differ in an *invisible* way must map
    to the same key, or an upsert/Modify silently creates a DUPLICATE event
    instead of updating the one already there. That is exactly the bug Sugit
    hit: two "Sarvasar Upanishad ~ 07" English records, identical on screen
    but differing by some mix of Unicode form, whitespace and which
    dash/tilde glyph was typed.

    The collapses applied — NFC, whitespace, dash/tilde family, case — never
    merge genuinely distinct talks: the series number (`~ 01` vs `~ 02`) is
    preserved, only the cosmetic surroundings are flattened. This key is used
    for comparison only; `events.title` still stores the title verbatim for
    display.
    """
    s = unicodedata.normalize("NFC", title or "")
    s = _TITLE_SEP_RE.sub("~", s)
    s = _TITLE_WS_RE.sub(" ", s).strip()
    return s.casefold()


# The "(English parts)"/"(Hindi parts)" qualifier is added to the *record title*
# during bilingual splits, but the discourse title inside the body text keeps the
# base name — so strip it before comparing a title to its body's first line.
_TITLE_PARTS_RE = re.compile(r"\s*\((?:english|hindi)\s+parts\)", re.IGNORECASE)


def _strip_parts(title: str) -> str:
    return _TITLE_PARTS_RE.sub("", title or "")


def _looks_like_title_line(text: str) -> bool:
    """True if `text` reads like a discourse-title line ("… ~ NN", short, Latin
    — the form OCTP bodies open with)."""
    s = (text or "").strip()
    return bool(s) and "~" in s and len(s) < 80 and sum(c.isascii() for c in s) > 0.8 * len(s)


def _title_content_warning(talk: "TalkRecord") -> str | None:
    """Catch a mislabeled header block — Sugit's 2026-07-02 incident, where a
    file named `Zen … _LHI.docx` carried `@title=Birthday Celebration 1978 ~ 01`
    on Zen content. That created a bogus record and left the real one stale.

    If the first body line is a discourse-title line naming a DIFFERENT talk than
    `@title` (ignoring the (parts) qualifier), return a warning string. Non-fatal
    — the operator decides — but it never again passes silently. Returns None
    when the title and body agree or the body doesn't open with a title line."""
    if not talk.paragraphs:
        return None
    first = talk.paragraphs[0].text.strip()
    if not _looks_like_title_line(first):
        return None
    if _canonical_title(_strip_parts(talk.title)) != _canonical_title(_strip_parts(first)):
        return (f"title/content mismatch — @title={talk.title!r} but the text "
                f"opens with {first!r}; this header may belong to a different discourse")
    return None


def _find_existing_event_id(conn: sqlite3.Connection, title: str, language: str) -> str | None:
    # Fast path: a byte-exact match hits the events(title) lookup and covers
    # the overwhelmingly common case (re-ingesting the same files unchanged).
    row = conn.execute(
        "SELECT id FROM events WHERE title = ? AND COALESCE(language, '') = ?",
        (title, language),
    ).fetchone()
    if row:
        return row[0]
    # Slow path, only when the exact match misses: tolerate invisible title
    # differences (Unicode form, whitespace, dash/tilde glyph) that would
    # otherwise make the upsert create a duplicate event. Scoped to the same
    # language and compared on the canonical key. This is an admin-only path
    # (ingest / batch-update), so the per-call candidate scan is acceptable;
    # the fast path keeps a clean bulk re-ingest off it entirely.
    want = _canonical_title(title)
    for rid, cand in conn.execute(
        "SELECT id, title FROM events WHERE COALESCE(language, '') = ?",
        (language,),
    ):
        if _canonical_title(cand or "") == want:
            return rid
    return None


def _delete_event_rows(conn: sqlite3.Connection, event_id: str) -> None:
    """Remove all paragraph + FTS + tag rows for an event. Used before re-inserting."""
    para_ids = [
        r[0]
        for r in conn.execute(
            "SELECT id FROM paragraphs WHERE event_id = ?", (event_id,)
        ).fetchall()
    ]
    if para_ids:
        placeholders = ",".join("?" * len(para_ids))
        conn.execute(
            f"DELETE FROM paragraphs_fts WHERE rowid IN ({placeholders})",
            para_ids,
        )
    conn.execute("DELETE FROM paragraphs WHERE event_id = ?", (event_id,))


def delete_record(
    conn: sqlite3.Connection, title: str, language: str
) -> tuple[str | None, int]:
    """Remove the (title, language) record entirely. Returns (event_id, paragraph_count)
    of what was deleted, or (None, 0) if no such record existed."""
    event_id = _find_existing_event_id(conn, title, language)
    if not event_id:
        return None, 0
    para_count = conn.execute(
        "SELECT COUNT(*) FROM paragraphs WHERE event_id = ?", (event_id,)
    ).fetchone()[0]
    _delete_event_rows(conn, event_id)
    if conn.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name='event_tags'"
    ).fetchone():
        conn.execute("DELETE FROM event_tags WHERE event_id = ?", (event_id,))
    conn.execute("DELETE FROM events WHERE id = ?", (event_id,))
    return event_id, para_count


def upsert(conn: sqlite3.Connection, talk: TalkRecord) -> tuple[str, bool]:
    """Insert or replace a talk in the DB. Returns (event_id, created_new)."""
    existing_id = _find_existing_event_id(conn, talk.title, talk.language)

    if existing_id:
        # Refresh metadata + paragraph rows. Keep the same event_id so
        # external references (bookmarks, analytics) stay stable.
        conn.execute(
            "UPDATE events SET date = ?, location = ?, language = ?, "
            "translated_from = ?, source_short = ? WHERE id = ?",
            (
                talk.time, talk.place, talk.language,
                talk.translated_from, talk.source_short, existing_id,
            ),
        )
        _delete_event_rows(conn, existing_id)
        event_id, created_new = existing_id, False
    else:
        event_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO events (id, title, date, location, language, "
            "translated_from, source_short) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (
                event_id, talk.title, talk.time, talk.place, talk.language,
                talk.translated_from, talk.source_short,
            ),
        )
        created_new = True

    conn.executemany(
        "INSERT INTO paragraphs (event_id, sequence_number, content, role, is_embedded) "
        "VALUES (?, ?, ?, ?, 0)",
        [
            (event_id, i, p.text, p.role)
            for i, p in enumerate(talk.paragraphs, start=1)
        ],
    )

    # Reinsert the FTS rows from the fresh paragraph rows, normalising
    # Devanagari to the same canonical form the rest of the index uses.
    norm_title = normalize_devanagari(talk.title)
    rows = conn.execute(
        "SELECT id, sequence_number, content FROM paragraphs WHERE event_id = ?",
        (event_id,),
    ).fetchall()
    conn.executemany(
        """
        INSERT INTO paragraphs_fts
            (content, title, event_id, paragraph_id, sequence_number, title_search)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        [
            (
                normalize_devanagari(content or ""),
                norm_title,
                event_id,
                pid,
                seq,
                norm_title,
            )
            for pid, seq, content in rows
        ],
    )

    # Optional @theme override — write as an event tag.
    if talk.theme:
        conn.execute(
            "CREATE TABLE IF NOT EXISTS event_tags ("
            "event_id TEXT, tag TEXT, PRIMARY KEY (event_id, tag))"
        )
        conn.execute(
            "INSERT OR IGNORE INTO event_tags (event_id, tag) VALUES (?, ?)",
            (event_id, talk.theme.lower()),
        )

    return event_id, created_new


# ─── CLI ────────────────────────────────────────────────────────────────────

def _walk(path: Path) -> list[Path]:
    if path.is_file():
        return [path] if path.suffix.lower() == ".docx" else []
    return sorted(p for p in path.rglob("*.docx") if not p.name.startswith("~$"))


def describe_no_docx(path: Path) -> str:
    """A clear, actionable message for the 'nothing to ingest' case — which used
    to pass as a silent success (Sugit's 2026-07-02 double-zip: the upload held a
    zip *inside* a zip, so 0 .docx were found and the run reported DONE having
    changed nothing). Flags a nested zip specifically."""
    nested = list(path.rglob("*.zip"))[:3] if path.is_dir() else []
    msg = f"No .docx files found under {path}."
    if nested:
        names = ", ".join(z.name for z in nested)
        msg += (f" It contains a zip *inside* it ({names}) — this looks"
                " double-zipped. Extract that inner zip (or upload a flat .zip of"
                " .docx files) and try again.")
    else:
        msg += (" Expected Word .docx files — check the archive isn't empty or in"
                " a different format.")
    return msg


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("path", type=Path, help="A .docx file, or a directory to walk")
    ap.add_argument("--dry-run", action="store_true", help="Parse but do not write to DB")
    ap.add_argument("--db", type=Path, default=DB_PATH, help=f"SQLite path (default {DB_PATH})")
    args = ap.parse_args(argv)

    if not args.path.exists():
        print(f"ERROR: {args.path} does not exist", file=sys.stderr)
        return 2

    files = _walk(args.path)
    if not files:
        # Loudly FAIL (exit 2), never silently succeed: a run that ingests
        # nothing almost always means a bad upload (double-zip, empty, wrong
        # format), and reporting success hid Sugit's whole 2026-07-02 incident.
        print(f"ERROR: {describe_no_docx(args.path)}", file=sys.stderr)
        return 2

    print(f"Found {len(files)} .docx file(s)")

    warn_count = 0
    if args.dry_run:
        for f in files:
            try:
                talk = parse_docx(f)
                print(
                    f"  OK   {f.name}: title={talk.title!r} "
                    f"language={talk.language} paragraphs={len(talk.paragraphs)}"
                )
                w = _title_content_warning(talk)
                if w:
                    warn_count += 1
                    print(f"  WARN {f.name}: {w}", file=sys.stderr)
            except Exception as ex:
                print(f"  FAIL {f.name}: {ex}")
        if warn_count:
            print(f"\n⚠ {warn_count} file(s) have a title/content mismatch — review before a real run.", file=sys.stderr)
        return 0

    if not args.db.exists():
        print(f"ERROR: DB not found at {args.db}", file=sys.stderr)
        return 2

    new_count = 0
    updated_count = 0
    failed_count = 0
    with sqlite3.connect(args.db) as conn:
        _ensure_translated_from_column(conn)
        _ensure_source_short_column(conn)
        _ensure_role_column(conn)
        for f in files:
            try:
                talk = parse_docx(f)
                w = _title_content_warning(talk)
                if w:
                    warn_count += 1
                    print(f"  WARN   {f.name}: {w}", file=sys.stderr)
                _, created_new = upsert(conn, talk)
                if created_new:
                    new_count += 1
                    print(f"  NEW    {talk.title}  [{talk.language}]  ({len(talk.paragraphs)} para)")
                else:
                    updated_count += 1
                    print(f"  UPDATE {talk.title}  [{talk.language}]  ({len(talk.paragraphs)} para)")
            except Exception as ex:
                failed_count += 1
                print(f"  FAIL   {f.name}: {ex}", file=sys.stderr)
        conn.commit()

    print(
        f"\nDone: {new_count} new, {updated_count} updated, {failed_count} failed"
        f"{f', {warn_count} warning(s)' if warn_count else ''}."
    )
    if warn_count:
        print(f"⚠ {warn_count} file(s) had a title/content mismatch — check them (a header may name the wrong discourse).", file=sys.stderr)
    if failed_count:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
